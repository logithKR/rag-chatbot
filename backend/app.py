#
# This file is: app.py (FULL CODE - FINAL VERSION)
#
import os
import sys
import atexit
import sqlite3
import shutil
from dotenv import load_dotenv
from datetime import datetime, date, timedelta
import pytz
import random
import hashlib
import gmail_send_otp

# --- NEW IMPORTS FOR READING DOCX/PDF ---
import docx  # From python-docx
import pypdf # From pypdf
# --- END NEW IMPORTS ---

# --- FLASK IMPORTS (with send_from_directory) ---
from flask import Flask, request, jsonify, send_from_directory, make_response

# --- CSV IMPORTS (for log download) ---
import io
import csv

from flask_cors import CORS

from langchain_community.document_loaders import DirectoryLoader
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

# --- 1. Load Environment Variables ---
load_dotenv()
if "GOOGLE_API_KEY" not in os.environ:
    print("Error: GOOGLE_API_KEY not found in .env file.")
    sys.exit(1)

EMAIL_SENDER = os.environ.get("email_id")
if not EMAIL_SENDER:
    print("Error: email_id not found in .env file.")
    print("This is required as the 'From' address for sending OTPs.")
    sys.exit(1)

# --- 2. Define All Paths ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENTS_PATH = os.path.join(BASE_DIR, "documents")
FAISS_INDEX_PATH = os.path.join(BASE_DIR, "faiss_index_google")
DATABASE_DIR = os.path.join(BASE_DIR, "database")

ESCALATED_DB_PATH = os.path.join(DATABASE_DIR, "user_details.db")
USERS_DB_PATH = os.path.join(DATABASE_DIR, "all_users.db")
OTP_DB_PATH = os.path.join(DATABASE_DIR, "otp.db")
# --- RENAMED: DB for Editor ---
EDITOR_STAFF_DB_PATH = os.path.join(DATABASE_DIR, "editor_staff.db")
EDIT_LOGS_DB_PATH = os.path.join(DATABASE_DIR, "edit_logs.db")

IST_TZ = pytz.timezone('Asia/Kolkata')

# --- 3. Initialize SQLite Databases ---
def setup_databases():
    """
    Creates the 'database' directory and initializes all database tables
    if they don't already exist.
    """
    try:
        os.makedirs(DATABASE_DIR, exist_ok=True)
        print(f"Database directory ensure_exists: {DATABASE_DIR}")

        # --- DB 1: Escalated User Details ---
        conn_details = sqlite3.connect(ESCALATED_DB_PATH)
        cursor_details = conn_details.cursor()
        cursor_details.execute("""
        CREATE TABLE IF NOT EXISTS escalated_queries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            user_name TEXT,
            email TEXT,
            phone_number TEXT,
            query_text TEXT NOT NULL,
            bot_response TEXT NOT NULL,
            status TEXT NOT NULL,
            remarks TEXT
        )
        """)
        conn_details.commit()
        conn_details.close()
        print(f"Escalated queries database initialized: {ESCALATED_DB_PATH}")

        # --- DB 2: All User Stats ---
        conn_users = sqlite3.connect(USERS_DB_PATH)
        cursor_users = conn_users.cursor()
        cursor_users.execute("""
        CREATE TABLE IF NOT EXISTS all_users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_name TEXT,
            email TEXT UNIQUE,
            phone_number TEXT,
            first_seen TEXT NOT NULL,
            last_seen TEXT NOT NULL
        )
        """)
        conn_users.commit()
        conn_users.close()
        print(f"All users database initialized: {USERS_DB_PATH}")

        # --- DB 3: OTP Requests ---
        conn_otp = sqlite3.connect(OTP_DB_PATH)
        cursor_otp = conn_otp.cursor()
        cursor_otp.execute("""
        CREATE TABLE IF NOT EXISTS otp_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT NOT NULL,
            otp_hash TEXT NOT NULL,
            expires_at TEXT NOT NULL
        )
        """)
        conn_otp.commit()
        conn_otp.close()
        print(f"OTP database initialized: {OTP_DB_PATH}")

        # --- DB 4: Editor Staff DB (File Renamed) ---
        conn_staff = sqlite3.connect(EDITOR_STAFF_DB_PATH)
        cursor_staff = conn_staff.cursor()
        
        # Table 1: Staff Members
        cursor_staff.execute("""
        CREATE TABLE IF NOT EXISTS staff_members (
            staff_id TEXT PRIMARY KEY,
            staff_name TEXT NOT NULL
        )
        """)
        try:
            cursor_staff.execute("INSERT INTO staff_members (staff_id, staff_name) VALUES (?, ?)", 
                                 ('BIT-STAFF-101', 'Dr. S. Ramesh'))
            conn_staff.commit()
            print("Default staff member 'BIT-STAFF-101' added.")
        except sqlite3.IntegrityError:
            print("Default staff member already exists.")
            
        # --- NEW: Table 2: Session Logs (with timestamps) ---
        cursor_staff.execute("""
        CREATE TABLE IF NOT EXISTS session_logs (
            session_id INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id TEXT NOT NULL,
            login_time TEXT NOT NULL,
            logout_time TEXT,
            FOREIGN KEY (staff_id) REFERENCES staff_members(staff_id)
        )
        """)
        print("Editor session logs table initialized.")
        # --- END NEW ---
        
        conn_staff.commit()
        conn_staff.close()
        print(f"Editor staff database initialized: {EDITOR_STAFF_DB_PATH}")

        # --- DB 5: Edit Logs ---
        conn_logs = sqlite3.connect(EDIT_LOGS_DB_PATH)
        cursor_logs = conn_logs.cursor()
        cursor_logs.execute("""
        CREATE TABLE IF NOT EXISTS edit_logs (
            log_id INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id TEXT NOT NULL,
            timestamp TEXT NOT NULL,
            action_performed TEXT NOT NULL,
            document_name TEXT NOT NULL,
            FOREIGN KEY (staff_id) REFERENCES staff_members(staff_id)
        )
        """)
        conn_logs.commit()
        conn_logs.close()
        print(f"Edit logs database initialized: {EDIT_LOGS_DB_PATH}")

    except Exception as e:
        print(f"Error initializing SQLite databases: {e}")
        sys.exit(1)

setup_databases()

# --- 4. Initialize Google Embedding Model ---
print("Initializing Google Generative AI Embedding model (uses API)...")
try:
    embedding_model = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    print("Google Embedding model (API) is ready.")
except Exception as e:
    print(f"Error initializing Google embedding model: {e}")
    sys.exit(1)

# --- 5. Setup the Vector Database (FAISS) ---
vectorstore = None
retriever = None

def create_new_faiss_index():
    print("Creating new FAISS index...")
    print(f"Loading documents from: {DOCUMENTS_PATH}")
    
    loader = DirectoryLoader(
        DOCUMENTS_PATH,
        glob="**/*.*",
        show_progress=True,
        use_multithreading=True,
        silent_errors=True
    )
    documents = loader.load()

    if not documents:
        print(f"No documents found in {DOCUMENTS_PATH}. Creating an empty index.")
        from langchain_core.documents import Document
        documents = [Document(page_content="No documents found.", metadata={"source": "dummy"})]
        print("Added a dummy document to create empty index.")

    print(f"Successfully loaded {len(documents)} documents.")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=300)
    texts = text_splitter.split_documents(documents)
    print(f"Split documents into {len(texts)} chunks.")

    print("Creating FAISS vector store... (This will call the Google API)")
    new_vectorstore = FAISS.from_documents(
        documents=texts,
        embedding=embedding_model
    )
    print(f"Saving FAISS index to: {FAISS_INDEX_PATH}")
    new_vectorstore.save_local(FAISS_INDEX_PATH)
    print("FAISS index created and saved successfully.")
    return new_vectorstore

def load_or_rebuild_index():
    global vectorstore, retriever
    print("Setting up the RAG chain...")
    
    if os.path.exists(FAISS_INDEX_PATH):
        print(f"Loading existing FAISS index from: {FAISS_INDEX_PATH}")
        try:
            vectorstore = FAISS.load_local(
                FAISS_INDEX_PATH,
                embedding_model,
                allow_dangerous_deserialization=True
            )
            print("FAISS index loaded successfully.")
        except Exception as e:
            print(f"Error loading FAISS index: {e}. Re-creating...")
            vectorstore = create_new_faiss_index()
    else:
        vectorstore = create_new_faiss_index()
    
    retriever = vectorstore.as_retriever(search_kwargs={"k": 30})
    print("Retriever is ready.")

def trigger_index_rebuild():
    global vectorstore, retriever
    print("--- TRIGGERING RAG INDEX REBUILD ---")
    
    if os.path.exists(FAISS_INDEX_PATH):
        try:
            shutil.rmtree(FAISS_INDEX_PATH)
            print(f"Deleted old index: {FAISS_INDEX_PATH}")
        except Exception as e:
            print(f"Error deleting old index: {e}")
            return False
            
    try:
        new_vectorstore = create_new_faiss_index()
        vectorstore = new_vectorstore
        retriever = vectorstore.as_retriever(search_kwargs={"k": 30})
        update_rag_chain()
        print("--- RAG INDEX REBUILD COMPLETE ---")
        return True
    except Exception as e:
        print(f"Error rebuilding index: {e}")
        return False

# --- 6. Initialize the RAG Chain ---
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.5)
template = """
You are **BIT AI Assistant**, the official virtual assistant of **Bannari Amman Institute of Technology (BIT)**.
Your primary goal is to be helpful, concise, and accurate. You must follow this strict logic flow:
────────────────────────────
 STEP 1: ANALYZE THE QUERY'S INTENT
────────────────────────────
First, determine the user's intent. Is it:
  A. General small talk (e.g., "Hello", "Who are you?", "How are you?")
  B. A question about BIT College (e.g., admissions, departments, fees, rules).
  C. A question about a topic completely unrelated to BIT (e.g., "Tell me about Harvard", "What is the weather?").

────────────────────────────
 STEP 2: CHOOSE YOUR RESPONSE STRATEGY
────────────────────────────
Based on your analysis in Step 1, you MUST follow one of these three paths:

**PATH A: If the intent is General Small Talk**
* Use your general knowledge (Gemini API) to respond naturally and conversationally.
* **DO NOT** use a static, repeated answer. Be friendly and flexible.
* (Example: If asked "Who are you?", reply: "I'm BIT AI Assistant, here to help you with questions about Bannari Amman Institute of Technology!")

**PATH B: If the intent is a BIT College Question**
* This is your most important job.
* Search the `{context}` retrieved from the FAISS vector store.

* **1. If the context CONTAINS the answer:**
    * **DO NOT** just copy the text from the context.
    * You **MUST** act as a summarizer. Read the context, understand the user's question, and then write a **short, clear, and concise summary**.
    * **Structure:** Use the "Formatting Rules" below to organize the data perfectly (Headings, Bullets, Bold Text).

* **2. If the context does NOT contain the answer:**
    * This means the information is completely missing from the BIT documents and must be escalated.
    * You **must** reply with this *exact* message and nothing else:
    "We have received your query, soon our concerned department will contact you. Thank You!"

**PATH C: If the intent is an Unrelated (non-BIT) Topic**
* This is for questions that are not small talk and not about BIT.
* You **must** reply with this *exact* message and nothing else:
  "I’m the official assistant of Bannari Amman Institute of Technology. Please ask questions related to BIT College."

────────────────────────────
 FORMATTING RULES (Strictly for PATH B):
────────────────────────────
* **Headings:** Use **Bold Headers** for main topics (e.g., **Eligibility Criteria**, **Tuition Fees**).
* **Summarize:** Be short, crisp, and to the point. No long blocks of text.
* **Keywords:** Use **bold** text for important numbers, dates, names, and deadlines.
* **Lists:** Use **• dot points** for features, lists, or requirements.
* **Sub-points:** If a topic has details, use nested bullets or short explanations under the main point.
* **Steps:** Use **1. Numbered steps** for procedures.
────────────────────────────
CONTEXT:
{context}

QUESTION:
{question}

ANSWER:
"""
prompt_template = PromptTemplate.from_template(template)

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

load_or_rebuild_index()

rag_chain = (
    {"context": retriever | format_docs, "question": RunnablePassthrough()}
    | prompt_template
    | llm
    | StrOutputParser()
)
print("RAG chain is ready.")

def update_rag_chain():
    global rag_chain
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt_template
        | llm
        | StrOutputParser()
    )
    print("RAG chain has been updated with new retriever.")

# --- 7. Create Flask App ---
app = Flask(__name__)
CORS(app)
print("Flask app created with CORS enabled.")

# --- 8. Helper Functions ---
def update_user_last_seen(email):
    try:
        timestamp = datetime.now(IST_TZ).isoformat()
        conn = sqlite3.connect(USERS_DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE all_users SET last_seen = ? WHERE email = ?",
            (timestamp, email)
        )
        conn.commit()
        conn.close()
        print(f"Updated last_seen for user: {email}")
    except Exception as e:
        print(f"Error updating last_seen for {email}: {e}")

def generate_otp(length=6):
    return "".join([str(random.randint(0, 9)) for _ in range(length)])

def hash_otp(otp):
    return hashlib.sha256(otp.encode()).hexdigest()

def log_editor_action(staff_id, action, document_name):
    try:
        timestamp = datetime.now(IST_TZ).isoformat()
        conn = sqlite3.connect(EDIT_LOGS_DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO edit_logs (staff_id, timestamp, action_performed, document_name)
            VALUES (?, ?, ?, ?)
            """,
            (staff_id, timestamp, action, document_name)
        )
        conn.commit()
        conn.close()
        print(f"Logged action: {action} for {document_name} by {staff_id}")
    except Exception as e:
        print(f"Error logging editor action: {e}")

# --- 9. User API Endpoints ---
@app.route('/request-otp', methods=['POST'])
def request_otp():
    try:
        data = request.json
        email = data.get('email')
        if not email:
            return jsonify({"error": "Email is required"}), 400
        otp = generate_otp()
        otp_hashed = hash_otp(otp)
        now = datetime.now(IST_TZ)
        expires = now + timedelta(minutes=10)
        expires_at_str = expires.isoformat()
        conn = sqlite3.connect(OTP_DB_PATH)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM otp_requests WHERE email = ?", (email,))
        cursor.execute(
            "INSERT INTO otp_requests (email, otp_hash, expires_at) VALUES (?, ?, ?)",
            (email, otp_hashed, expires_at_str)
        )
        conn.commit()
        conn.close()
        if not gmail_send_otp.send_otp_email_gmail(email, otp):
            print(f"Failed to send OTP email to {email}")
            return jsonify({"error": "Failed to send OTP email"}), 500
        return jsonify({"status": "success", "message": "OTP sent to your email."})
    except Exception as e:
        print(f"Error in /request-otp: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/login', methods=['POST'])
def login_user():
    try:
        data = request.json
        user_name = data.get('name')
        email = data.get('email')
        phone = data.get('phone')
        otp_submitted = data.get('otp')
        if not email or not otp_submitted:
            return jsonify({"error": "Email and OTP are required"}), 400
        
        conn_otp = sqlite3.connect(OTP_DB_PATH)
        cursor_otp = conn_otp.cursor()
        now_str = datetime.now(IST_TZ).isoformat()
        otp_submitted_hash = hash_otp(otp_submitted)
        cursor_otp.execute(
            "SELECT id FROM otp_requests WHERE email = ? AND otp_hash = ? AND expires_at > ?",
            (email, otp_submitted_hash, now_str)
        )
        valid_otp_row = cursor_otp.fetchone()
        
        if not valid_otp_row:
            conn_otp.close()
            return jsonify({"error": "Invalid or expired OTP"}), 401
        
        cursor_otp.execute("DELETE FROM otp_requests WHERE id = ?", (valid_otp_row[0],))
        conn_otp.commit()
        conn_otp.close()
        
        timestamp = datetime.now(IST_TZ).isoformat()
        conn = sqlite3.connect(USERS_DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM all_users WHERE email = ?", (email,))
        user = cursor.fetchone()
        
        if user:
            cursor.execute(
                "UPDATE all_users SET last_seen = ? WHERE email = ?",
                (timestamp, email)
            )
            message = "Login successful"
        else:
            cursor.execute(
                "INSERT INTO all_users (user_name, email, phone_number, first_seen, last_seen) VALUES (?, ?, ?, ?, ?)",
                (user_name, email, phone, timestamp, timestamp)
            )
            message = "Registration successful"
        
        conn.commit()
        conn.close()
        return jsonify({"status": "success", "message": message, "email": email})

    except sqlite3.IntegrityError:
        if 'conn_otp' in locals() and conn_otp: conn_otp.close()
        if 'conn' in locals() and conn: conn.close()
        update_user_last_seen(email)
        return jsonify({"status": "success", "message": "Login successful", "email": email})
    except Exception as e:
        if 'conn_otp' in locals() and conn_otp: conn_otp.close()
        if 'conn' in locals() and conn: conn.close()
        print(f"Error in /login: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        query = data.get('query')
        user_name = data.get('user_name')
        email = data.get('email')
        phone_number = data.get('phone_number')

        if not query or not email:
            return jsonify({"error": "Query and email are required"}), 400

        update_user_last_seen(email)
        print(f"\nReceived query from {email}: {query}")
        
        answer = rag_chain.invoke(query)
        print(f"Generated answer: {answer}")

        trigger_message = "We have received your query, soon our concerned department will contact you. Thank You!"
        if answer.strip() == trigger_message:
            print("Trigger message detected! Saving to escalated_queries.db...")
            try:
                timestamp = datetime.now(IST_TZ).isoformat()
                conn_details = sqlite3.connect(ESCALATED_DB_PATH)
                cursor_details = conn_details.cursor()
                cursor_details.execute(
                    "INSERT INTO escalated_queries (timestamp, user_name, email, phone_number, query_text, bot_response, status, remarks) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (timestamp, user_name or 'Not Provided', email or 'Not Provided', phone_number or 'Not Provided', query, answer, "Initiated", "")
                )
                conn_details.commit()
                conn_details.close()
                print("Escalated query saved successfully.")
            except Exception as e:
                print(f"Error logging to escalated_queries.db: {e}")
        else:
            print("Standard response. Not saving to escalated database.")
        return jsonify({"answer": answer})
    except Exception as e:
        print(f"An error occurred in the /chat route: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

# --- 10. Admin Endpoints ---
@app.route('/admin/stats', methods=['GET'])
def get_admin_stats():
    try:
        conn_users = sqlite3.connect(USERS_DB_PATH)
        cursor_users = conn_users.cursor()
        cursor_users.execute("SELECT COUNT(id) FROM all_users")
        total_unique_users = cursor_users.fetchone()[0]
        today_str = str(datetime.now(IST_TZ).date())
        cursor_users.execute("SELECT COUNT(id) FROM all_users WHERE last_seen LIKE ?", (today_str + '%',))
        today_users = cursor_users.fetchone()[0]
        conn_users.close()
        conn_details = sqlite3.connect(ESCALATED_DB_PATH)
        cursor_details = conn_details.cursor()
        cursor_details.execute("SELECT COUNT(id) FROM escalated_queries WHERE status = 'Initiated'")
        total_escalated = cursor_details.fetchone()[0]
        cursor_details.execute("SELECT COUNT(id) FROM escalated_queries WHERE status = 'Finished'")
        total_solved = cursor_details.fetchone()[0]
        conn_details.close()
        stats = {
            "totalUniqueUsers": total_unique_users, "todayUsers": today_users,
            "totalEscalated": total_escalated, "totalSolved": total_solved
        }
        return jsonify(stats)
    except Exception as e:
        print(f"Error in /admin/stats: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/escalated-queries', methods=['GET'])
def get_escalated_queries():
    try:
        conn = sqlite3.connect(ESCALATED_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM escalated_queries ORDER BY timestamp DESC")
        rows = cursor.fetchall()
        queries = [dict(row) for row in rows]
        conn.close()
        return jsonify(queries)
    except Exception as e:
        print(f"Error in /admin/escalated-queries: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/all-users', methods=['GET'])
def get_all_users():
    try:
        conn = sqlite3.connect(USERS_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM all_users ORDER BY last_seen DESC")
        rows = cursor.fetchall()
        users = [dict(row) for row in rows]
        conn.close()
        return jsonify(users)
    except Exception as e:
        print(f"Error in /admin/all-users: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/update-query/<int:query_id>', methods=['PUT'])
def update_escalated_query(query_id):
    try:
        data = request.json
        status = data.get('status')
        remarks = data.get('remarks')
        if not status or remarks is None:
            return jsonify({"error": "Status and remarks are required"}), 400
        conn = sqlite3.connect(ESCALATED_DB_PATH)
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE escalated_queries SET status = ?, remarks = ? WHERE id = ?",
            (status, remarks, query_id)
        )
        conn.commit()
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({"error": "Query not found"}), 404
        conn.close()
        return jsonify({"status": "success", "message": "Query updated"})
    except Exception as e:
        print(f"Error in /admin/update-query: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/escalated-users', methods=['GET'])
def get_escalated_users():
    try:
        conn = sqlite3.connect(ESCALATED_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("""
            SELECT user_name, email, phone_number, COUNT(*) as query_count
            FROM escalated_queries WHERE status = 'Initiated'
            GROUP BY email ORDER BY query_count DESC, user_name ASC
        """)
        rows = cursor.fetchall()
        users = [dict(row) for row in rows]
        conn.close()
        return jsonify(users)
    except Exception as e:
        print(f"Error in /admin/escalated-users: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/user-queries/<email>', methods=['GET'])
def get_user_queries(email):
    try:
        conn = sqlite3.connect(ESCALATED_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM escalated_queries WHERE email = ? ORDER BY timestamp DESC", (email,))
        rows = cursor.fetchall()
        queries = [dict(row) for row in rows]
        conn.close()
        return jsonify(queries)
    except Exception as e:
        print(f"Error in /admin/user-queries: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/solved-users', methods=['GET'])
def get_solved_users():
    try:
        conn = sqlite3.connect(ESCALATED_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("""
            SELECT user_name, email, phone_number, COUNT(*) as query_count
            FROM escalated_queries WHERE status = 'Finished'
            GROUP BY email ORDER BY query_count DESC, user_name ASC
        """)
        rows = cursor.fetchall()
        users = [dict(row) for row in rows]
        conn.close()
        return jsonify(users)
    except Exception as e:
        print(f"Error in /admin/solved-users: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/today-users', methods=['GET'])
def get_today_users():
    try:
        today_str = str(datetime.now(IST_TZ).date())
        conn = sqlite3.connect(USERS_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM all_users WHERE last_seen LIKE ? ORDER BY last_seen DESC", (today_str + '%',))
        rows = cursor.fetchall()
        users = [dict(row) for row in rows]
        conn.close()
        return jsonify(users)
    except Exception as e:
        print(f"Error in /admin/today-users: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500
        
# --- 11. Admin-Editor Endpoints (New) ---
@app.route('/admin/editor-staff', methods=['GET'])
def get_editor_staff():
    """ Fetches all staff members and their last login time. """
    try:
        conn = sqlite3.connect(EDITOR_STAFF_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        # Get staff and their most recent login time from session_logs
        cursor.execute("""
            SELECT 
                s.staff_id, 
                s.staff_name, 
                MAX(sl.login_time) as last_login
            FROM staff_members s
            LEFT JOIN session_logs sl ON s.staff_id = sl.staff_id
            GROUP BY s.staff_id, s.staff_name
            ORDER BY s.staff_name ASC
        """)
        rows = cursor.fetchall()
        staff = [dict(row) for row in rows]
        
        conn.close()
        return jsonify(staff)
    except Exception as e:
        print(f"Error in /admin/editor-staff: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/editor-logs', methods=['GET'])
def get_editor_logs():
    """ Fetches editor logs with filters for search and date. """
    try:
        # Get query params
        search_term = request.args.get('search', '')
        filter_period = request.args.get('filter', 'all')

        conn_logs = sqlite3.connect(EDIT_LOGS_DB_PATH)
        conn_logs.row_factory = sqlite3.Row
        cursor_logs = conn_logs.cursor()
        
        # Attach the staff database to join tables
        cursor_logs.execute("ATTACH DATABASE ? AS editor_staff", (EDITOR_STAFF_DB_PATH,))
        
        query = """
            SELECT 
                el.log_id,
                el.timestamp,
                el.staff_id,
                es.staff_name,
                el.action_performed,
                el.document_name
            FROM main.edit_logs el
            LEFT JOIN editor_staff.staff_members es ON el.staff_id = es.staff_id
            WHERE (IFNULL(es.staff_name, '') LIKE ? OR el.staff_id LIKE ?)
        """
        
        params = [f'%{search_term}%', f'%{search_term}%']
        
        # Add date filter logic
        if filter_period == '10days':
            cutoff_date = (datetime.now(IST_TZ) - timedelta(days=10)).isoformat()
            query += " AND el.timestamp >= ?"
            params.append(cutoff_date)
        elif filter_period == '30days':
            cutoff_date = (datetime.now(IST_TZ) - timedelta(days=30)).isoformat()
            query += " AND el.timestamp >= ?"
            params.append(cutoff_date)
        
        query += " ORDER BY el.timestamp DESC"
        
        cursor_logs.execute(query, tuple(params))
        rows = cursor_logs.fetchall()
        logs = [dict(row) for row in rows]
        
        cursor_logs.execute("DETACH DATABASE editor_staff")
        conn_logs.close()
        
        return jsonify(logs)
    except Exception as e:
        print(f"Error in /admin/editor-logs: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/admin/download-editor-logs', methods=['GET'])
def download_editor_logs():
    """ Fetches editor logs with filters and returns as CSV. """
    try:
        search_term = request.args.get('search', '')
        filter_period = request.args.get('filter', 'all')

        conn_logs = sqlite3.connect(EDIT_LOGS_DB_PATH)
        conn_logs.row_factory = sqlite3.Row
        cursor_logs = conn_logs.cursor()
        
        cursor_logs.execute("ATTACH DATABASE ? AS editor_staff", (EDITOR_STAFF_DB_PATH,))
        
        query = """
            SELECT 
                el.log_id,
                el.timestamp,
                el.staff_id,
                es.staff_name,
                el.action_performed,
                el.document_name
            FROM main.edit_logs el
            LEFT JOIN editor_staff.staff_members es ON el.staff_id = es.staff_id
            WHERE (IFNULL(es.staff_name, '') LIKE ? OR el.staff_id LIKE ?)
        """
        params = [f'%{search_term}%', f'%{search_term}%']

        if filter_period == '10days':
            cutoff_date = (datetime.now(IST_TZ) - timedelta(days=10)).isoformat()
            query += " AND el.timestamp >= ?"
            params.append(cutoff_date)
        elif filter_period == '30days':
            cutoff_date = (datetime.now(IST_TZ) - timedelta(days=30)).isoformat()
            query += " AND el.timestamp >= ?"
            params.append(cutoff_date)
        
        query += " ORDER BY el.timestamp DESC"
        
        cursor_logs.execute(query, tuple(params))
        rows = cursor_logs.fetchall()
        
        cursor_logs.execute("DETACH DATABASE editor_staff")
        conn_logs.close()

        # Create CSV in memory
        si = io.StringIO()
        cw = csv.writer(si)
        
        if rows:
            cw.writerow(rows[0].keys())
            for row in rows:
                cw.writerow(row)
        else:
            cw.writerow(['log_id', 'timestamp', 'staff_id', 'staff_name', 'action_performed', 'document_name'])

        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = "attachment; filename=editor_logs.csv"
        output.headers["Content-type"] = "text/csv"
        return output

    except Exception as e:
        print(f"Error in /admin/download-editor-logs: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500


# --- 12. Editor Endpoints ---
@app.route('/editor/login', methods=['POST'])
def editor_login():
    """ Handles the first-step (static) editor login. """
    try:
        data = request.json
        email = data.get('email')
        password = data.get('password')
        if (email == 'editor.bitra@bitsathy.ac.in' and password == 'editor.pass.bitra@12345'):
            return jsonify({"status": "success", "message": "Editor login successful"})
        else:
            return jsonify({"error": "Invalid editor credentials"}), 401
    except Exception as e:
        print(f"Error in /editor/login: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/verify-staff', methods=['POST'])
def verify_staff():
    """ 
    Handles 'Get or Create' for staff members.
    Creates a session and returns a session_id.
    """
    try:
        data = request.json
        staff_id = data.get('staff_id')
        staff_name = data.get('staff_name')

        if not staff_id or not staff_name:
            return jsonify({"error": "Staff ID and Name are required"}), 400

        conn = sqlite3.connect(EDITOR_STAFF_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM staff_members WHERE staff_id = ?", (staff_id,))
        staff = cursor.fetchone()
        
        if staff:
            print(f"Verified existing staff: {staff_id}")
        else:
            print(f"Creating new staff member: {staff_id} - {staff_name}")
            cursor.execute("INSERT INTO staff_members (staff_id, staff_name) VALUES (?, ?)", (staff_id, staff_name))
            conn.commit()

        login_time = datetime.now(IST_TZ).isoformat()
        cursor.execute(
            "INSERT INTO session_logs (staff_id, login_time) VALUES (?, ?)",
            (staff_id, login_time)
        )
        conn.commit()
        
        session_id = cursor.lastrowid
        print(f"Created new session {session_id} for {staff_id}")

        conn.close()
        return jsonify({"status": "success", "message": "Staff verified", "staff_id": staff_id, "session_id": session_id})
    
    except sqlite3.IntegrityError as ie:
        print(f"Integrity Error: {ie}")
        conn.close()
        return jsonify({"error": "A database error occurred."}), 500
    except Exception as e:
        if conn: conn.close()
        print(f"Error in /editor/verify-staff: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/logout', methods=['POST'])
def editor_logout():
    """ Handles editor logout by updating the session log. """
    try:
        data = request.json
        session_id = data.get('session_id')
        if not session_id:
            return jsonify({"error": "Session ID is required"}), 400

        conn = sqlite3.connect(EDITOR_STAFF_DB_PATH)
        cursor = conn.cursor()
        
        logout_time = datetime.now(IST_TZ).isoformat()
        cursor.execute(
            "UPDATE session_logs SET logout_time = ? WHERE session_id = ?",
            (logout_time, session_id)
        )
        conn.commit()
        
        print(f"Logged out session {session_id}")
        conn.close()
        return jsonify({"status": "success", "message": "Logout successful"})
        
    except Exception as e:
        if conn: conn.close()
        print(f"Error in /editor/logout: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/get-folders', methods=['GET'])
def get_folders():
    """ Returns a list of all subdirectories in DOCUMENTS_PATH. """
    try:
        folders = ["Main Folder"]
        for item in os.listdir(DOCUMENTS_PATH):
            item_path = os.path.join(DOCUMENTS_PATH, item)
            if os.path.isdir(item_path):
                folders.append(item)
        
        return jsonify(folders)
    except Exception as e:
        print(f"Error in /editor/get-folders: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/get-documents', methods=['GET'])
def get_documents():
    """ Lists all files in the DOCUMENTS_PATH, including subfolders. """
    try:
        files = []
        for root, _, filenames in os.walk(DOCUMENTS_PATH):
            for filename in filenames:
                relative_path = os.path.relpath(os.path.join(root, filename), DOCUMENTS_PATH)
                files.append(relative_path.replace("\\", "/"))
        return jsonify(files)
    except Exception as e:
        print(f"Error in /editor/get-documents: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/download-file', methods=['GET'])
def download_file():
    """ Downloads a specific file from the documents folder. """
    try:
        filename = request.args.get('filename')
        if not filename:
            return jsonify({"error": "Filename is required"}), 400
        if '..' in filename or os.path.isabs(filename):
            return jsonify({"error": "Invalid filename"}), 400
        
        file_path = os.path.join(DOCUMENTS_PATH, filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
        
        # We need to find the directory and the actual filename
        # e.g., "departments/file.txt" -> dir="departments", name="file.txt"
        directory = os.path.dirname(filename)
        name = os.path.basename(filename)
        
        return send_from_directory(
            os.path.join(DOCUMENTS_PATH, directory), 
            name, 
            as_attachment=True
        )
    except Exception as e:
        print(f"Error in /editor/download-file: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/get-document-content', methods=['POST'])
def get_document_content():
    """ Fetches the text content of a document, extracting from docx/pdf. """
    try:
        data = request.json
        filename = data.get('filename')
        if not filename or '..' in filename:
            return jsonify({"error": "Invalid filename"}), 400
        
        file_path = os.path.join(DOCUMENTS_PATH, filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        content = ""
        try:
            if filename.endswith('.txt') or filename.endswith('.md'):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif filename.endswith('.docx'):
                document = docx.Document(file_path)
                full_text = [para.text for para in document.paragraphs]
                content = "\n".join(full_text)
            elif filename.endswith('.pdf'):
                reader = pypdf.PdfReader(file_path)
                full_text = [page.extract_text() for page in reader.pages]
                content = "\n".join(full_text)
            else:
                content = f"--- This is a binary file ({filename}) and cannot be edited. ---"
        except Exception as read_error:
            print(f"Error reading {filename}: {read_error}")
            content = f"--- Error reading file {filename}. File may be corrupt. ---"

        return jsonify({"filename": filename, "content": content})
    except Exception as e:
        print(f"Error in /editor/get-document-content: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/update-document', methods=['POST'])
def update_document():
    """ Updates an existing .txt or .md document. Does NOT rebuild index. """
    try:
        data = request.json
        filename = data.get('filename')
        content = data.get('content')
        staff_id = data.get('staff_id')

        if not filename or content is None or not staff_id:
            return jsonify({"error": "Filename, content, and staff_id are required"}), 400
        
        if not (filename.endswith('.txt') or filename.endswith('.md')):
            print(f"BLOCKED save for non-text file: {filename}")
            return jsonify({"error": f"Editing is not allowed for {filename} files."}), 403

        file_path = os.path.join(DOCUMENTS_PATH, filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        log_editor_action(staff_id, "Document Edited", filename)
        return jsonify({"status": "success", "message": f"'{filename}' updated."})
    except Exception as e:
        print(f"Error in /editor/update-document: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/add-document', methods=['POST'])
def add_document():
    """ Adds a new .txt or .md document. Does NOT rebuild index. """
    try:
        data = request.json
        filename = data.get('filename')
        content = data.get('content')
        staff_id = data.get('staff_id')
        folder = data.get('folder')

        if not filename or content is None or not staff_id or folder is None:
            return jsonify({"error": "All fields are required"}), 400
            
        if not (filename.endswith('.txt') or filename.endswith('.md')):
             filename += ".txt"
        
        if '..' in filename:
            return jsonify({"error": "Invalid filename"}), 400
        
        target_folder = DOCUMENTS_PATH
        if folder != "Main Folder":
            target_folder = os.path.join(DOCUMENTS_PATH, folder)

        file_path = os.path.join(target_folder, filename)
        
        if os.path.exists(file_path):
            return jsonify({"error": "File with this name already exists in this folder"}), 409
            
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        log_filename = filename if folder == "Main Folder" else f"{folder}/{filename}"
        log_editor_action(staff_id, "Document Added", log_filename)
        return jsonify({"status": "success", "message": f"'{log_filename}' added."})
    except Exception as e:
        print(f"Error in /editor/add-document: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/upload-file', methods=['POST'])
def upload_file():
    """ Uploads any file, replacing if it exists. Does NOT rebuild index. """
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        
        file = request.files['file']
        folder = request.form.get('folder')
        staff_id = request.form.get('staff_id')

        if not file or not folder or not staff_id:
            return jsonify({"error": "Missing file, folder, or staff_id"}), 400

        filename = file.filename
        if '..' in filename:
             return jsonify({"error": "Invalid filename"}), 400

        target_folder = DOCUMENTS_PATH
        if folder != "Main Folder":
            target_folder = os.path.join(DOCUMENTS_PATH, folder)
        
        file_path = os.path.join(target_folder, filename)
        
        os.makedirs(target_folder, exist_ok=True)
        file.save(file_path)
        
        log_filename = filename if folder == "Main Folder" else f"{folder}/{filename}"
        log_editor_action(staff_id, "File Uploaded", log_filename)
        
        return jsonify({"status": "success", "message": f"'{log_filename}' uploaded."})
    except Exception as e:
        print(f"Error in /editor/upload-file: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/delete-document', methods=['POST'])
def delete_document():
    """ Deletes a document. Does NOT rebuild index. """
    try:
        data = request.json
        filename = data.get('filename')
        staff_id = data.get('staff_id')

        if not filename or not staff_id:
            return jsonify({"error": "Filename and staff_id are required"}), 400
        
        if '..' in filename:
            return jsonify({"error": "Invalid filename"}), 400
            
        file_path = os.path.join(DOCUMENTS_PATH, filename)
        
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        os.remove(file_path)
        log_editor_action(staff_id, "Document Deleted", filename)
        return jsonify({"status": "success", "message": f"'{filename}' deleted."})
    except Exception as e:
        print(f"Error in /editor/delete-document: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

@app.route('/editor/commit-index', methods=['POST'])
def commit_index():
    """ Rebuilds the RAG index. """
    try:
        data = request.json
        staff_id = data.get('staff_id')
        if not staff_id:
             return jsonify({"error": "Staff ID is required"}), 400
             
        rebuild_success = trigger_index_rebuild()
        if not rebuild_success:
            return jsonify({"error": "Failed to rebuild RAG index"}), 500
        
        log_editor_action(staff_id, "RAG Index Rebuilt", "N/A")
        return jsonify({"status": "success", "message": "RAG index has been rebuilt."})
    except Exception as e:
        print(f"Error in /editor/commit-index: {e}")
        return jsonify({"error": "An internal server error occurred"}), 500

# --- 13. Run the Flask App ---
if __name__ == "__main__":
    print("\n--- BIT Chatbot Backend is Starting ---")
    print(f"--- Database File: {EDITOR_STAFF_DB_PATH} ---")
    print("Using Gmail API for sending OTPs.")
    print("Editor and Admin endpoints are active.")
    print("PDF and DOCX text extraction is ENABLED.")
    print("Access the API at http://127.0.0.1:5000")
    app.run(debug=True, port=5000)